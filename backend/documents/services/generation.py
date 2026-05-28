"""Генерация docx (python-docx) и PDF (reportlab) для документов."""

from __future__ import annotations

import os
import platform
import secrets
from pathlib import Path

from django.conf import settings
from django.utils import timezone

from documents.models import Document
from documents.services.document_layout import build_document_context, format_money


def _output_dir() -> Path:
    path = Path(settings.MEDIA_ROOT) / "documents"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _setup_docx_styles(doc):
    from docx.shared import Pt

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for level in range(1, 4):
        if f"Heading {level}" in doc.styles:
            h = doc.styles[f"Heading {level}"]
            h.font.name = "Times New Roman"
            h.font.bold = True


def _add_docx_paragraph(doc, text, *, bold=False, align_center=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    return p


def _generate_docx(document: Document, ctx: dict, out_path: Path) -> None:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm

    doc = DocxDocument()
    _setup_docx_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    # Заголовок по центру (как в типовых формах)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run(f"{ctx['title']}")
    t_run.bold = True
    t_run.font.name = "Times New Roman"
    t_run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run(f"№ {ctx['number']} от {ctx['date']}")
    s_run.font.name = "Times New Roman"
    s_run.font.size = Pt(12)

    doc.add_paragraph()

    seller = ctx["seller"]
    buyer = ctx["buyer"]
    _add_docx_paragraph(doc, "Продавец:", bold=True)
    _add_docx_paragraph(doc, f"{seller['name']}")
    _add_docx_paragraph(doc, f"ИНН {seller['inn']}, КПП {seller['kpp']}")
    _add_docx_paragraph(doc, f"Адрес: {seller['address']}")
    _add_docx_paragraph(doc, f"Тел.: {seller['phone']}")
    doc.add_paragraph()
    _add_docx_paragraph(doc, "Покупатель:", bold=True)
    _add_docx_paragraph(doc, f"{buyer['name']}")
    _add_docx_paragraph(doc, f"ИНН {buyer['inn']}")
    _add_docx_paragraph(doc, f"Адрес: {buyer['address']}")
    doc.add_paragraph()
    _add_docx_paragraph(doc, f"Основание: заказ № {ctx['order_number']}")
    if ctx["client_name"]:
        _add_docx_paragraph(doc, f"Клиент: {ctx['client_name']}")
    if ctx["manager_name"]:
        _add_docx_paragraph(doc, f"Менеджер: {ctx['manager_name']}")
    doc.add_paragraph()

    # Таблица позиций
    if ctx["items"]:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["№ п/п", "Наименование", "Ед. изм.", "Кол-во", "Цена, руб.", "Сумма, руб."]
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for p in hdr_cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)

        for row in ctx["items"]:
            cells = table.add_row().cells
            values = [
                str(row["num"]),
                row["name"],
                row["unit"],
                str(row["qty"]),
                format_money(row["price"]),
                format_money(row["amount"]),
            ]
            for cell, val in zip(cells, values):
                cell.text = val
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(11)

        doc.add_paragraph()
        _add_docx_paragraph(doc, f"Итого: {ctx['total_formatted']} руб.", bold=True)
        _add_docx_paragraph(doc, ctx["vat_note"])
    else:
        _add_docx_paragraph(doc, "Позиции: отсутствуют")

    doc.add_paragraph()

    if ctx["doc_type"] == "кп":
        _add_docx_paragraph(doc, "Срок действия предложения: 10 (десять) календарных дней.")
        doc.add_paragraph()

    if ctx["doc_type"] == "акт":
        _add_docx_paragraph(
            doc,
            "Вышеперечисленные работы (услуги) выполнены полностью и в срок. "
            "Заказчик претензий по объёму, качеству и срокам оказания услуг не имеет.",
        )
        doc.add_paragraph()

    if ctx["responsible_name"]:
        _add_docx_paragraph(doc, f"Ответственное лицо (исполнитель): {ctx['responsible_name']}")
    if ctx["created_by_name"]:
        _add_docx_paragraph(doc, f"Составил: {ctx['created_by_name']}")
    sign_line = ctx["signed_by_name"] or "________________"
    _add_docx_paragraph(doc, f"Подпись / ЭЦП: {sign_line}")
    _add_docx_paragraph(doc, f"Статус документа: {ctx['status']}")

    doc.save(out_path)


def _pdf_font_paths() -> tuple[str | None, str | None]:
    """Обычный и жирный TTF для кириллицы."""
    regular = bold = None
    candidates_reg = []
    candidates_bold = []
    if platform.system() == "Windows":
        fonts_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
        candidates_reg = [
            fonts_dir / "times.ttf",
            fonts_dir / "Times.ttf",
            fonts_dir / "arial.ttf",
        ]
        candidates_bold = [
            fonts_dir / "timesbd.ttf",
            fonts_dir / "Timesbd.ttf",
            fonts_dir / "arialbd.ttf",
        ]
    candidates_reg.extend(
        [
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/TTF/DejaVuSans.ttf"),
        ]
    )
    candidates_bold.extend(
        [
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            Path("/usr/share/fonts/TTF/DejaVuSans-Bold.ttf"),
        ]
    )
    for p in candidates_reg:
        if p.is_file():
            regular = str(p)
            break
    for p in candidates_bold:
        if p.is_file():
            bold = str(p)
            break
    return regular, bold or regular


def _register_pdf_fonts() -> tuple[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    reg_name, bold_name = "DocTimes", "DocTimes-Bold"
    if reg_name not in pdfmetrics.getRegisteredFontNames():
        reg_path, bold_path = _pdf_font_paths()
        if reg_path:
            pdfmetrics.registerFont(TTFont(reg_name, reg_path))
        if bold_path:
            pdfmetrics.registerFont(TTFont(bold_name, bold_path))
        else:
            bold_name = reg_name
    return reg_name, bold_name


def _generate_pdf(document: Document, ctx: dict, out_path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    reg_font, bold_font = _register_pdf_fonts()
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "BaseRu",
        parent=styles["Normal"],
        fontName=reg_font,
        fontSize=12,
        leading=14,
        alignment=TA_LEFT,
    )
    title_style = ParagraphStyle(
        "TitleRu",
        parent=base,
        fontName=bold_font,
        fontSize=14,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    center_style = ParagraphStyle("CenterRu", parent=base, alignment=TA_CENTER)

    doc_pdf = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    story = []
    story.append(Paragraph(ctx["title"], title_style))
    story.append(Paragraph(f"№ {ctx['number']} от {ctx['date']}", center_style))
    story.append(Spacer(1, 12))

    seller = ctx["seller"]
    buyer = ctx["buyer"]
    blocks = [
        f"<b>Продавец:</b><br/>{seller['name']}<br/>"
        f"ИНН {seller['inn']}, КПП {seller['kpp']}<br/>"
        f"Адрес: {seller['address']}<br/>Тел.: {seller['phone']}",
        f"<b>Покупатель:</b><br/>{buyer['name']}<br/>"
        f"ИНН {buyer['inn']}<br/>Адрес: {buyer['address']}",
        f"Основание: заказ № {ctx['order_number']}<br/>"
        f"Клиент: {ctx['client_name'] or '—'}"
        + (f"<br/>Менеджер: {ctx['manager_name']}" if ctx["manager_name"] else ""),
    ]
    for block in blocks:
        story.append(Paragraph(block, base))
        story.append(Spacer(1, 8))

    if ctx["items"]:
        table_data = [
            ["№", "Наименование", "Ед.", "Кол-во", "Цена", "Сумма"],
        ]
        for row in ctx["items"]:
            table_data.append(
                [
                    str(row["num"]),
                    row["name"][:40],
                    row["unit"],
                    str(row["qty"]),
                    format_money(row["price"]),
                    format_money(row["amount"]),
                ]
            )
        table_data.append(["", "", "", "", "Итого:", f"{ctx['total_formatted']} руб."])

        col_widths = [1 * cm, 7 * cm, 1.5 * cm, 2 * cm, 2.5 * cm, 3 * cm]
        t = Table(table_data, colWidths=col_widths, repeatRows=1)
        t.setStyle(
            TableStyle(
                [
                    ("FONT", (0, 0), (-1, -1), reg_font, 10),
                    ("FONT", (0, 0), (-1, 0), bold_font, 10),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("ALIGN", (1, 1), (1, -2), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("FONT", (4, -1), (5, -1), bold_font, 10),
                ]
            )
        )
        story.append(t)
        story.append(Spacer(1, 8))
        story.append(Paragraph(ctx["vat_note"], base))

    if ctx["doc_type"] == "кп":
        story.append(Spacer(1, 6))
        story.append(Paragraph("Срок действия предложения: 10 (десять) календарных дней.", base))

    if ctx["doc_type"] == "акт":
        story.append(Spacer(1, 6))
        story.append(
            Paragraph(
                "Вышеперечисленные работы (услуги) выполнены полностью и в срок. "
                "Заказчик претензий по объёму, качеству и срокам оказания услуг не имеет.",
                base,
            )
        )

    story.append(Spacer(1, 16))
    if ctx["responsible_name"]:
        story.append(Paragraph(f"<b>Ответственное лицо:</b> {ctx['responsible_name']}", base))
    if ctx["created_by_name"]:
        story.append(Paragraph(f"<b>Составил:</b> {ctx['created_by_name']}", base))
    sign_line = ctx["signed_by_name"] or "________________"
    story.append(Paragraph(f"<b>Подпись / ЭЦП:</b> {sign_line}", base))
    story.append(Paragraph(f"Статус документа: {ctx['status']}", base))

    doc_pdf.build(story)


def generate_document_files(document: Document) -> Document:
    """Создаёт docx и pdf, сохраняет относительные пути в модели."""
    document = (
        Document.objects.select_related(
            "order",
            "order__client",
            "order__manager",
            "created_by",
            "responsible",
            "signed_by",
        )
        .prefetch_related("order__items__product")
        .get(pk=document.pk)
    )
    ctx = build_document_context(document)
    out_dir = _output_dir()
    suffix = secrets.token_hex(4)
    docx_name = f"{document.number}_{suffix}.docx"
    pdf_name = f"{document.number}_{suffix}.pdf"
    docx_path = out_dir / docx_name
    pdf_path = out_dir / pdf_name

    _generate_docx(document, ctx, docx_path)
    _generate_pdf(document, ctx, pdf_path)

    document.docx_file = f"documents/{docx_name}"
    document.pdf_file = f"documents/{pdf_name}"
    document.save(update_fields=["docx_file", "pdf_file", "updated_at"])
    return document
